"""
Archive Search v10.2
====================

What this version provides
--------------------------
- A Tkinter desktop interface for searching one or two root folders.
- Read-only extraction from Excel, Word, and PDF files where supported.
- Per-root JSON caching so repeat searches are much faster.
- A text-based results pane with clickable paths.

Safety model
------------
The program is designed to avoid modifying indexed source documents. It writes
only its own cache file inside each indexed root folder and prefers read-only
open operations wherever platform support is available.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, Iterator, List, Optional, Sequence, Tuple

import datetime as dt
import json
import math
import os
import queue
import shlex
import subprocess
import sys
import tempfile
import threading
import time

import openpyxl
import xlrd

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    DocxDocument = None
    HAS_DOCX = False

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    PdfReader = None
    HAS_PDF = False

import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    import pythoncom
    import win32com.client
    HAS_WIN32_COM = True
except ImportError:
    HAS_WIN32_COM = False


# =============================================================================
# Configuration
# =============================================================================
DEFAULT_PRIMARY_SEARCH_FOLDER = Path(r"C:\Primary\Folder")
DEFAULT_SECONDARY_SEARCH_FOLDER = Path(r"C:\Secondary\Folder")

OPENPYXL_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
XLRD_EXTENSIONS = {".xls"}
WORKBOOK_EXTENSIONS = OPENPYXL_EXTENSIONS | XLRD_EXTENSIONS

WORD_XML_EXTENSIONS = {".docx", ".docm", ".dotx", ".dotm"}
WORD_LEGACY_EXTENSIONS = {".doc"}
WORD_EXTENSIONS = WORD_XML_EXTENSIONS | WORD_LEGACY_EXTENSIONS

PDF_EXTENSIONS = {".pdf"}
TEXT_DOCUMENT_EXTENSIONS = WORD_EXTENSIONS | PDF_EXTENSIONS

# The cache lives inside each indexed root so every root is self-contained and can
# be moved independently without a central application data folder.
CACHE_FILE_NAME = ".Archive_Search_cache.json"
CACHE_SCHEMA_VERSION = 2
CACHE_FILE_PERMISSIONS = 0o600

STATUS_UPDATE_INTERVAL_SECONDS = 0.35
OUTPUT_BATCH_BLOCKS = 20

HIGHLIGHT_TAG_NAME = "match_highlight"
HIGHLIGHT_BACKGROUND = "#fff59d"
HIGHLIGHT_FOREGROUND = "#000000"


# =============================================================================
# Small data structures
# =============================================================================
@dataclass(frozen=True)
class SearchRoot:
    """One search root configured in the UI."""

    label: str
    path: Path
    include_subfolders: bool


@dataclass(frozen=True)
class FileRecord:
    """A discovered source file under a specific search root."""

    root_path: Path
    file_path: Path

    @property
    def relative_path(self) -> str:
        """
        Return this file path relative to the root it was discovered under.

        Relative paths are used as stable cache keys so the same file can be
        found again even when the absolute root folder differs between systems.
        """
        return self.file_path.relative_to(self.root_path).as_posix()


@dataclass(frozen=True)
class SearchResult:
    """A single result rendered in the output panel."""

    file_type: str
    document_name: str
    path: str
    sheet: Optional[str]
    row_number: Optional[int]
    row_text: str
    line_number: Optional[int]
    line_before: str
    matched_line: str
    line_after: str
    match_count: Optional[int] = None


@dataclass(frozen=True)
class DiscoveryResult:
    """
    The result of file-tree discovery for one root.

    Only the values needed by later stages are stored here. Folder reuse/rescan
    counts are reported immediately via status messages, so they do not need to be
    carried through the rest of the pipeline.
    """

    file_records: Tuple[FileRecord, ...]
    manifest_directories: Dict[str, Dict[str, object]]
    used_manifest: bool


# =============================================================================
# Generic helpers
# =============================================================================
def get_app_folder() -> Path:
    """
    Return the folder that contains the running program.

    When the script is packaged as an executable, this resolves to the executable
    folder. During normal Python execution it resolves to the folder that contains
    this source file.
    """
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def compact_whitespace(value) -> str:
    """
    Collapse repeated whitespace and convert the incoming value to a clean string.

    This keeps cached content and search previews predictable even when source
    files contain extra spaces, tabs, or line breaks.
    """
    if value is None:
        return ""
    return " ".join(str(value).split())


def normalize_text(value) -> str:
    """
    Convert incoming values into a lower-case search-friendly representation.

    Dates and times are normalised to a consistent human-readable format before the
    text is lowered so searches behave the same across spreadsheets and documents.
    """
    if value is None:
        return ""

    if isinstance(value, dt.datetime):
        return value.strftime("%d/%m/%Y").lower()
    if isinstance(value, dt.date):
        return value.strftime("%d/%m/%Y").lower()
    if isinstance(value, dt.time):
        return value.strftime("%H:%M").lower()

    return compact_whitespace(value).lower()


def display_text(value) -> str:
    """
    Convert a raw cell or document value into a display-safe string.

    This preserves readable formatting for dates, times, and whole-number floats so
    search results look sensible in the user interface and in cached output.
    """
    if value is None:
        return ""

    if isinstance(value, float):
        if math.isfinite(value) and value.is_integer():
            return str(int(value))
        return str(value)

    if isinstance(value, dt.datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, dt.date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, dt.time):
        return value.strftime("%H:%M")

    return compact_whitespace(value)


def build_highlight_terms(search_terms: Sequence[str]) -> List[str]:
    """
    Build a de-duplicated list of highlight terms from the current search input.

    The full term is kept first, then the individual words are added in descending
    length order so longer matches are highlighted before shorter fragments.
    """
    seen = set()
    items: List[str] = []

    for term in sorted(search_terms, key=len, reverse=True):
        whole = term.strip().lower()
        if whole and whole not in seen:
            items.append(whole)
            seen.add(whole)

        parts = [part.strip().lower() for part in whole.split() if part.strip()]
        for part in sorted(parts, key=len, reverse=True):
            if part and part not in seen:
                items.append(part)
                seen.add(part)

    return items


def text_matches_search(text: str, search_terms: Sequence[str], match_mode: str = "any") -> bool:
    """
    Return True when the provided text satisfies the current search mode.

    In 'any' mode a single term is enough. In 'all' mode every parsed term must be
    present in the normalised text.
    """
    normalized = normalize_text(text)
    if not normalized:
        return False

    if match_mode == "all":
        return all(term in normalized for term in search_terms)
    return any(term in normalized for term in search_terms)


def format_error_block(message: str) -> str:
    """
    Format an extraction error so it can be displayed in the UI output cleanly.
    """
    return f"{message}\n\n"


def build_excel_result_payload(file_path: Path, sheet: str, row_number: int, row_text: str) -> SearchResult:
    """
    Build the result payload used by the text-output UI for one Excel row match.
    """
    return SearchResult(
        file_type="excel",
        document_name=file_path.name,
        path=str(file_path),
        sheet=sheet,
        row_number=row_number,
        row_text=row_text,
        line_number=None,
        line_before="",
        matched_line="",
        line_after="",
    )


def build_text_result_payload(
    file_path: Path,
    file_type: str,
    line_number: int,
    line_before: str,
    matched_line: str,
    line_after: str,
    match_count: Optional[int] = None,
) -> SearchResult:
    """
    Build the result payload used by the text-output UI for one document match.
    """
    return SearchResult(
        file_type=file_type,
        document_name=file_path.name,
        path=str(file_path),
        sheet=None,
        row_number=None,
        row_text="",
        line_number=line_number,
        line_before=compact_whitespace(line_before),
        matched_line=compact_whitespace(matched_line),
        line_after=compact_whitespace(line_after),
        match_count=match_count,
    )


def pluralize(count: int, singular: str, plural: Optional[str] = None) -> str:
    """
    Return a singular or plural label that matches the supplied count.
    """
    return singular if count == 1 else (plural or singular + "s")


def should_ignore_filename(filename: str) -> bool:
    """
    Exclude temporary Office lock files and this program's own cache files.
    """
    if filename.startswith("~$"):
        return True
    if filename == CACHE_FILE_NAME or filename.startswith(CACHE_FILE_NAME + "."):
        return True
    return False


def directory_key(relative_directory: Path) -> str:
    """
    Convert a relative directory path into the normalised manifest key format.
    """
    key = relative_directory.as_posix()
    return "." if key in {"", "."} else key


# =============================================================================
# Value conversion for legacy .xls
# =============================================================================
def convert_xls_value(value, ctype, datemode):
    """
    Convert an xlrd cell value into a more useful Python value.

    Legacy XLS files expose raw value/ctype pairs, so this helper restores dates,
    integers, booleans, and blank cells into forms used consistently elsewhere.
    """
    if ctype == xlrd.XL_CELL_DATE:
        try:
            return xlrd.xldate.xldate_as_datetime(value, datemode)
        except Exception:
            return value

    if ctype == xlrd.XL_CELL_NUMBER:
        if float(value).is_integer():
            return int(value)
        return value

    if ctype == xlrd.XL_CELL_BOOLEAN:
        return bool(value)

    if ctype in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK):
        return None

    return value


# =============================================================================
# Read-only extraction layer
# =============================================================================
class ReadOnlyWordSession:
    """Shared read-only Microsoft Word session for legacy .doc extraction."""

    def __init__(self) -> None:
        """
        Initialise the wrapper with no active Word COM application yet.
        """
        self.word = None

    def __enter__(self) -> "ReadOnlyWordSession":
        """
        Start a hidden Microsoft Word instance configured for read-only automation.
        """
        if not HAS_WIN32_COM:
            raise RuntimeError("pywin32 is not installed.")

        pythoncom.CoInitialize()
        self.word = win32com.client.DispatchEx("Word.Application")
        self.word.Visible = False
        self.word.DisplayAlerts = 0

        try:
            self.word.AutomationSecurity = 3
        except Exception:
            pass

        try:
            self.word.ScreenUpdating = False
        except Exception:
            pass

        return self

    def __exit__(self, exc_type, exc, tb) -> None:
        """
        Close the shared Word session and clean up COM state safely.
        """
        if self.word is not None:
            try:
                self.word.Quit()
            except Exception:
                pass

        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

    def extract_lines(self, file_path: Path) -> List[str]:
        """
        Open a legacy Word document read-only and return a cleaned list of text lines.
        """
        if self.word is None:
            raise RuntimeError("Word session is not open.")

        document = None
        try:
            document = self.word.Documents.Open(
                str(file_path.resolve()),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
                Revert=False,
                OpenAndRepair=False,
                NoEncodingDialog=True,
            )
            raw_text = document.Content.Text or ""
            raw_text = raw_text.replace("\r", "\n").replace("\x07", " ")
            lines = [compact_whitespace(line) for line in raw_text.splitlines()]
            return [line for line in lines if line]
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass


class DocumentExtractor:
    """Read-only extraction for every supported file type."""

    @staticmethod
    def infer_file_type(suffix: str) -> str:
        """
        Map a file suffix to the internal file-type label used by the cache and UI.
        """
        suffix = suffix.lower()
        if suffix in WORKBOOK_EXTENSIONS:
            return "excel"
        if suffix in WORD_EXTENSIONS:
            return "word"
        if suffix in PDF_EXTENSIONS:
            return "pdf"
        return "unknown"

    @staticmethod
    def build_row_text(row_values: Sequence[object]) -> Optional[str]:
        """
        Join the non-empty values from one spreadsheet row into a searchable string.
        """
        display_cells = [display_text(value) for value in row_values]
        row_text = " | ".join(cell for cell in display_cells if cell).strip(" |")
        return row_text or None

    def extract_xlsx_rows(self, file_path: Path) -> List[Dict[str, object]]:
        """
        Read workbook rows from modern Excel files using openpyxl in read-only mode.
        """
        wb = openpyxl.load_workbook(
            file_path,
            read_only=True,
            data_only=True,
            keep_links=False,
        )
        try:
            rows: List[Dict[str, object]] = []
            for ws in wb.worksheets:
                for row_number, row in enumerate(ws.iter_rows(values_only=True), start=1):
                    row_text = self.build_row_text(row)
                    if not row_text:
                        continue
                    rows.append({
                        "sheet": ws.title,
                        "row_number": row_number,
                        "row_text": row_text,
                    })
            return rows
        finally:
            wb.close()

    def extract_xls_rows(self, file_path: Path) -> List[Dict[str, object]]:
        """
        Read workbook rows from legacy XLS files using xlrd.
        """
        book = xlrd.open_workbook(file_path, on_demand=True)
        try:
            datemode = book.datemode
            rows: List[Dict[str, object]] = []
            for sheet_name in book.sheet_names():
                sheet = book.sheet_by_name(sheet_name)
                for row_idx in range(sheet.nrows):
                    raw_values = sheet.row_values(row_idx)
                    raw_types = sheet.row_types(row_idx)
                    row_values = [
                        convert_xls_value(value, ctype, datemode)
                        for value, ctype in zip(raw_values, raw_types)
                    ]
                    row_text = self.build_row_text(row_values)
                    if not row_text:
                        continue
                    rows.append({
                        "sheet": sheet_name,
                        "row_number": row_idx + 1,
                        "row_text": row_text,
                    })
            return rows
        finally:
            try:
                book.release_resources()
            except Exception:
                pass

    def extract_docx_lines(self, file_path: Path) -> List[str]:
        """
        Extract paragraph and table text from modern Word files.
        """
        if not HAS_DOCX:
            raise RuntimeError("python-docx is not installed.")

        document = DocxDocument(str(file_path))
        lines: List[str] = []

        for paragraph in document.paragraphs:
            text = compact_whitespace(paragraph.text)
            if text:
                lines.append(text)

        for table in document.tables:
            for row in table.rows:
                cell_values = [compact_whitespace(cell.text) for cell in row.cells]
                row_text = " | ".join(value for value in cell_values if value)
                if row_text:
                    lines.append(row_text)

        return lines

    def extract_word_lines(self, file_path: Path, word_session: Optional[ReadOnlyWordSession] = None) -> List[str]:
        """
        Choose the appropriate Word extraction path for DOCX/DOC-style files.
        """
        suffix = file_path.suffix.lower()

        if suffix in WORD_LEGACY_EXTENSIONS:
            if word_session is not None:
                return word_session.extract_lines(file_path)
            if not HAS_WIN32_COM:
                raise RuntimeError(".doc files require pywin32 and Microsoft Word.")
            with ReadOnlyWordSession() as session:
                return session.extract_lines(file_path)

        if HAS_DOCX:
            return self.extract_docx_lines(file_path)

        if HAS_WIN32_COM:
            if word_session is not None:
                return word_session.extract_lines(file_path)
            with ReadOnlyWordSession() as session:
                return session.extract_lines(file_path)

        raise RuntimeError(".docx files require python-docx or pywin32 + Microsoft Word.")

    def extract_pdf_lines(self, file_path: Path) -> List[str]:
        """
        Extract text lines from every page in a PDF file.
        """
        if not HAS_PDF:
            raise RuntimeError("pypdf is not installed.")

        reader = PdfReader(str(file_path))
        lines: List[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            page_lines = [compact_whitespace(line) for line in page_text.splitlines()]
            lines.extend(line for line in page_lines if line)
        return lines

    def build_error_record(self, file_path: Path, size, mtime_ns, message: str) -> Dict[str, object]:
        """
        Build a cache record that preserves file metadata alongside an extraction error.
        """
        return {
            "size": size,
            "mtime_ns": mtime_ns,
            "suffix": file_path.suffix.lower(),
            "file_type": self.infer_file_type(file_path.suffix),
            "rows": [],
            "lines": [],
            "error": message,
        }

    def build_cache_record(
        self,
        file_path: Path,
        size: Optional[int],
        mtime_ns: Optional[int],
        word_session: Optional[ReadOnlyWordSession] = None,
    ) -> Dict[str, object]:
        """
        Extract content for one file and package it into the cache record structure.
        """
        suffix = file_path.suffix.lower()

        try:
            if suffix in OPENPYXL_EXTENSIONS:
                return {
                    "size": size,
                    "mtime_ns": mtime_ns,
                    "suffix": suffix,
                    "file_type": "excel",
                    "rows": self.extract_xlsx_rows(file_path),
                    "lines": [],
                    "error": None,
                }

            if suffix in XLRD_EXTENSIONS:
                return {
                    "size": size,
                    "mtime_ns": mtime_ns,
                    "suffix": suffix,
                    "file_type": "excel",
                    "rows": self.extract_xls_rows(file_path),
                    "lines": [],
                    "error": None,
                }

            if suffix in WORD_EXTENSIONS:
                return {
                    "size": size,
                    "mtime_ns": mtime_ns,
                    "suffix": suffix,
                    "file_type": "word",
                    "rows": [],
                    "lines": self.extract_word_lines(file_path, word_session=word_session),
                    "error": None,
                }

            if suffix in PDF_EXTENSIONS:
                return {
                    "size": size,
                    "mtime_ns": mtime_ns,
                    "suffix": suffix,
                    "file_type": "pdf",
                    "rows": [],
                    "lines": self.extract_pdf_lines(file_path),
                    "error": None,
                }

            return self.build_error_record(file_path, size, mtime_ns, f"Unsupported file type: {suffix}")
        except Exception as exc:
            return self.build_error_record(file_path, size, mtime_ns, f"Could not open {file_path}: {exc}")


# =============================================================================
# Cache manager + manifest-assisted discovery
# =============================================================================
class CacheManager:
    """Handles discovery, cache refresh, cache writing, and cached searching."""

    def __init__(self, extractor: DocumentExtractor) -> None:
        """
        Store the extractor used for all file refresh operations.
        """
        self.extractor = extractor

    @staticmethod
    def cache_path_for_root(root_path: Path) -> Path:
        """
        Return the JSON cache path that belongs to a configured search root.
        """
        return root_path / CACHE_FILE_NAME

    @staticmethod
    def record_matches_file(record: Dict[str, object], size: int, mtime_ns: int, suffix: str) -> bool:
        """
        Return True when a cached record still matches the current file on disk.
        """
        return (
            isinstance(record, dict)
            and record.get("size") == size
            and record.get("mtime_ns") == mtime_ns
            and record.get("suffix") == suffix.lower()
        )

    @staticmethod
    def new_cache_document(root: SearchRoot, allowed_extensions: Sequence[str]) -> Dict[str, object]:
        """
        Create a fresh cache document skeleton for one search root.
        """
        return {
            "schema_version": CACHE_SCHEMA_VERSION,
            "root_path": str(root.path),
            "include_subfolders": bool(root.include_subfolders),
            "allowed_extensions": sorted({ext.lower() for ext in allowed_extensions}),
            "generated_at_utc": dt.datetime.now(dt.timezone.utc).isoformat(),
            "files": {},
            "manifest": {"directories": {}},
        }

    def load_cache(self, root_path: Path) -> Optional[Dict[str, object]]:
        """
        Load and validate a previously written cache document for a root folder.
        """
        cache_path = self.cache_path_for_root(root_path)
        if not cache_path.exists():
            return None

        try:
            with cache_path.open("r", encoding="utf-8") as handle:
                payload = json.load(handle)
            if not isinstance(payload, dict):
                return None
            if payload.get("schema_version") != CACHE_SCHEMA_VERSION:
                return None
            if not isinstance(payload.get("files"), dict):
                return None
            manifest = payload.get("manifest", {})
            if not isinstance(manifest, dict):
                return None
            directories = manifest.get("directories", {})
            if not isinstance(directories, dict):
                return None
            return payload
        except Exception:
            return None

    def write_cache(self, root_path: Path, cache_document: Dict[str, object]) -> None:
        """
        Write the cache atomically so partial writes do not corrupt the cache file.
        """
        cache_path = self.cache_path_for_root(root_path)
        temp_fd, temp_name = tempfile.mkstemp(
            prefix=CACHE_FILE_NAME + ".",
            suffix=".tmp",
            dir=str(root_path),
        )
        try:
            with os.fdopen(temp_fd, "w", encoding="utf-8", newline="\n") as handle:
                json.dump(cache_document, handle, ensure_ascii=False, separators=(",", ":"))
            os.replace(temp_name, cache_path)
            try:
                os.chmod(cache_path, CACHE_FILE_PERMISSIONS)
            except Exception:
                pass
        finally:
            if os.path.exists(temp_name):
                try:
                    os.remove(temp_name)
                except Exception:
                    pass

    def _cache_is_compatible(
        self,
        previous_cache: Optional[Dict[str, object]],
        root: SearchRoot,
        allowed_extensions: Sequence[str],
    ) -> bool:
        """
        Check whether an existing cache can be reused for the current root settings.
        """
        if previous_cache is None:
            return False
        if previous_cache.get("root_path") != str(root.path):
            return False
        if previous_cache.get("include_subfolders") != root.include_subfolders:
            return False
        cached_extensions = sorted({ext.lower() for ext in previous_cache.get("allowed_extensions", [])})
        current_extensions = sorted({ext.lower() for ext in allowed_extensions})
        return cached_extensions == current_extensions

    def discover_root_files(
        self,
        root: SearchRoot,
        allowed_extensions: Sequence[str],
        cancel_event: threading.Event,
        status_callback: Callable[[str], None],
        previous_cache: Optional[Dict[str, object]],
    ) -> DiscoveryResult:
        """Discover files using either a full scan or the cached directory manifest."""
        if not self._cache_is_compatible(previous_cache, root, allowed_extensions):
            return self._full_discover_root(root, allowed_extensions, cancel_event, status_callback)
        return self._manifest_discover_root(root, allowed_extensions, cancel_event, status_callback, previous_cache)

    def _full_discover_root(
        self,
        root: SearchRoot,
        allowed_extensions: Sequence[str],
        cancel_event: threading.Event,
        status_callback: Callable[[str], None],
    ) -> DiscoveryResult:
        """
        Walk the configured folder tree from scratch and collect supported files.

        The discovery manifest intentionally stores only directory mtimes and child
        folder names. File names are revalidated from the cache's file list, so
        storing them here would just enlarge the cache without helping later runs.
        """
        allowed = {ext.lower() for ext in allowed_extensions}
        file_records: List[FileRecord] = []
        manifest_dirs: Dict[str, Dict[str, object]] = {}
        last_status_time = 0.0
        total_found = 0
        rescanned_dirs = 0

        def scan_directory(directory_path: Path, relative_directory: Path) -> None:
            # Walk this folder once, collecting supported files and a lightweight
            # manifest entry that can speed up future discovery passes.
            nonlocal last_status_time, total_found, rescanned_dirs
            if cancel_event.is_set():
                return

            key = directory_key(relative_directory)
            rescanned_dirs += 1

            try:
                dir_stat = directory_path.stat()
                entries = sorted(os.scandir(directory_path), key=lambda entry: entry.name.lower())
            except OSError:
                return

            subdirs: List[str] = []

            for entry in entries:
                if cancel_event.is_set():
                    return

                entry_name = entry.name
                entry_path = directory_path / entry_name

                try:
                    is_symlink = entry.is_symlink()
                except OSError:
                    continue
                if is_symlink:
                    continue

                try:
                    if entry.is_dir(follow_symlinks=False):
                        subdirs.append(entry_name)
                        continue
                except OSError:
                    continue

                try:
                    is_file = entry.is_file(follow_symlinks=False)
                except OSError:
                    continue
                if not is_file:
                    continue

                if should_ignore_filename(entry_name):
                    continue
                if entry_path.suffix.lower() not in allowed:
                    continue

                file_records.append(FileRecord(root.path, entry_path))
                total_found += 1

                now = time.monotonic()
                if (now - last_status_time) >= STATUS_UPDATE_INTERVAL_SECONDS:
                    status_callback(f"Discovering {root.label.lower()} files... found {total_found}")
                    last_status_time = now

            manifest_dirs[key] = {
                "mtime_ns": getattr(dir_stat, "st_mtime_ns", int(dir_stat.st_mtime * 1_000_000_000)),
                "subdirs": subdirs,
            }

            if root.include_subfolders:
                for subdir_name in subdirs:
                    scan_directory(directory_path / subdir_name, relative_directory / subdir_name)

        scan_directory(root.path, Path("."))
        file_records.sort(key=lambda item: str(item.file_path).lower())
        return DiscoveryResult(
            file_records=tuple(file_records),
            manifest_directories=manifest_dirs,
            used_manifest=False,
        )

    def _manifest_discover_root(
        self,
        root: SearchRoot,
        allowed_extensions: Sequence[str],
        cancel_event: threading.Event,
        status_callback: Callable[[str], None],
        previous_cache: Dict[str, object],
    ) -> DiscoveryResult:
        """
        Discover files using the cached file list + cached directory manifest.

        Strategy:
        1. Re-check every cached file directly by path. This catches modifications,
           removals, extension changes, and symlink issues without a full walk.
        2. Re-scan only directories whose directory mtime has changed. Those are
           the places where new/removed/renamed files are likely to appear.
        """
        allowed = {ext.lower() for ext in allowed_extensions}
        previous_files: Dict[str, Dict[str, object]] = previous_cache.get("files", {})
        previous_dirs: Dict[str, Dict[str, object]] = previous_cache.get("manifest", {}).get("directories", {})

        file_records: List[FileRecord] = []
        manifest_dirs: Dict[str, Dict[str, object]] = {}
        known_rel_paths: set[str] = set()
        last_status_time = 0.0
        reused_dirs = 0
        rescanned_dirs = 0

        # Pass 1: directly validate all previously known files.
        # This cheaply confirms whether cached paths still exist without needing
        # to enumerate every directory in the tree again.
        total_known = len(previous_files)
        for index, relative_path in enumerate(sorted(previous_files.keys()), start=1):
            if cancel_event.is_set():
                return DiscoveryResult(tuple(file_records), manifest_dirs, True)

            file_path = root.path / Path(relative_path)
            try:
                if file_path.is_symlink() or not file_path.is_file():
                    continue
            except OSError:
                continue

            if should_ignore_filename(file_path.name):
                continue
            if file_path.suffix.lower() not in allowed:
                continue

            try:
                file_path.stat()
            except OSError:
                continue

            file_records.append(FileRecord(root.path, file_path))
            known_rel_paths.add(relative_path)

            now = time.monotonic()
            if total_known and ((now - last_status_time) >= STATUS_UPDATE_INTERVAL_SECONDS or index == total_known):
                status_callback(f"Checking known files in {root.label.lower()} cache... {index}/{total_known}")
                last_status_time = now

        # Pass 2: re-scan only changed directories to discover new/removed names.
        # Directory modification times act as a coarse but fast signal that the
        # file listing beneath that folder may need to be rebuilt.
        def scan_directory(directory_path: Path, relative_directory: Path) -> None:
            # Walk this folder once, collecting supported files and a lightweight
            # manifest entry that can speed up future discovery passes.
            nonlocal reused_dirs, rescanned_dirs, last_status_time
            if cancel_event.is_set():
                return

            key = directory_key(relative_directory)
            previous_dir = previous_dirs.get(key)

            try:
                dir_stat = directory_path.stat()
            except OSError:
                return

            current_dir_mtime_ns = getattr(dir_stat, "st_mtime_ns", int(dir_stat.st_mtime * 1_000_000_000))

            if previous_dir and previous_dir.get("mtime_ns") == current_dir_mtime_ns:
                # The directory entry list appears unchanged. Reuse cached knowledge of
                # its child directories and avoid enumerating this directory again.
                reused_dirs += 1
                manifest_dirs[key] = previous_dir
                if root.include_subfolders:
                    for subdir_name in previous_dir.get("subdirs", []):
                        scan_directory(directory_path / subdir_name, relative_directory / subdir_name)
                return

            rescanned_dirs += 1
            try:
                entries = sorted(os.scandir(directory_path), key=lambda entry: entry.name.lower())
            except OSError:
                return

            subdirs: List[str] = []

            for entry in entries:
                if cancel_event.is_set():
                    return

                entry_name = entry.name
                entry_path = directory_path / entry_name

                try:
                    is_symlink = entry.is_symlink()
                except OSError:
                    continue
                if is_symlink:
                    continue

                try:
                    if entry.is_dir(follow_symlinks=False):
                        subdirs.append(entry_name)
                        continue
                except OSError:
                    continue

                try:
                    is_file = entry.is_file(follow_symlinks=False)
                except OSError:
                    continue
                if not is_file:
                    continue

                if should_ignore_filename(entry_name):
                    continue
                if entry_path.suffix.lower() not in allowed:
                    continue

                relative_file = directory_key(relative_directory / entry_name)
                if relative_file not in known_rel_paths:
                    file_records.append(FileRecord(root.path, entry_path))
                    known_rel_paths.add(relative_file)

            manifest_dirs[key] = {
                "mtime_ns": current_dir_mtime_ns,
                "subdirs": subdirs,
            }

            now = time.monotonic()
            if (now - last_status_time) >= STATUS_UPDATE_INTERVAL_SECONDS:
                status_callback(
                    f"Checking {root.label.lower()} folders... reused {reused_dirs}, rescanned {rescanned_dirs}"
                )
                last_status_time = now

            if root.include_subfolders:
                for subdir_name in subdirs:
                    scan_directory(directory_path / subdir_name, relative_directory / subdir_name)

        scan_directory(root.path, Path("."))
        file_records.sort(key=lambda item: str(item.file_path).lower())
        status_callback(
            f"{root.label}: discovery reused {reused_dirs} folder {pluralize(reused_dirs, 'entry', 'entries')} and rescanned {rescanned_dirs}"
        )
        return DiscoveryResult(
            file_records=tuple(file_records),
            manifest_directories=manifest_dirs,
            used_manifest=True,
        )

    def refresh_root_cache(
        self,
        root: SearchRoot,
        discovery_result: DiscoveryResult,
        allowed_extensions: Sequence[str],
        previous_cache: Optional[Dict[str, object]],
        cancel_event: threading.Event,
        status_callback: Callable[[str], None],
    ) -> Dict[str, object]:
        """
        Reuse unchanged cache entries and refresh only files whose metadata changed.

        The expensive step in this program is content extraction, so this method
        aggressively reuses previous records whenever size/mtime/suffix still match.
        """
        previous_files: Dict[str, Dict[str, object]] = (previous_cache or {}).get("files", {})
        current_files: Dict[str, Dict[str, object]] = {}

        refreshed_count = 0
        reused_count = 0
        legacy_doc_records: List[Tuple[str, Path, int, int]] = []
        file_records = discovery_result.file_records
        total_files = len(file_records)
        last_status_time = 0.0

        for index, record in enumerate(file_records, start=1):
            if cancel_event.is_set():
                return previous_cache or self.new_cache_document(root, allowed_extensions)

            file_path = record.file_path
            rel_path = record.relative_path
            try:
                stat = file_path.stat()
            except OSError as exc:
                current_files[rel_path] = self.extractor.build_error_record(
                    file_path=file_path,
                    size=None,
                    mtime_ns=None,
                    message=f"Could not stat {file_path}: {exc}",
                )
                refreshed_count += 1
                continue

            old_record = previous_files.get(rel_path)
            suffix = file_path.suffix.lower()
            # When the file metadata matches the cached record exactly we can reuse
            # the extracted content and skip a potentially expensive re-read.
            if old_record and self.record_matches_file(old_record, stat.st_size, stat.st_mtime_ns, suffix):
                current_files[rel_path] = old_record
                reused_count += 1
            else:
                refreshed_count += 1
                # Legacy .doc extraction is batched so a single shared Word session
                # can service multiple files instead of launching Word repeatedly.
                if suffix in WORD_LEGACY_EXTENSIONS and HAS_WIN32_COM:
                    legacy_doc_records.append((rel_path, file_path, stat.st_size, stat.st_mtime_ns))
                else:
                    current_files[rel_path] = self.extractor.build_cache_record(
                        file_path=file_path,
                        size=stat.st_size,
                        mtime_ns=stat.st_mtime_ns,
                    )

            now = time.monotonic()
            if total_files and (index == total_files or (now - last_status_time) >= STATUS_UPDATE_INTERVAL_SECONDS):
                status_callback(f"{root.label}: checking content cache {index}/{total_files}")
                last_status_time = now

        if legacy_doc_records:
            try:
                with ReadOnlyWordSession() as word_session:
                    total_legacy = len(legacy_doc_records)
                    last_status_time = 0.0
                    for index, (rel_path, file_path, size, mtime_ns) in enumerate(legacy_doc_records, start=1):
                        if cancel_event.is_set():
                            return previous_cache or self.new_cache_document(root, allowed_extensions)

                        current_files[rel_path] = self.extractor.build_cache_record(
                            file_path=file_path,
                            size=size,
                            mtime_ns=mtime_ns,
                            word_session=word_session,
                        )

                        now = time.monotonic()
                        if index == total_legacy or (now - last_status_time) >= STATUS_UPDATE_INTERVAL_SECONDS:
                            status_callback(f"{root.label}: refreshing legacy .doc cache {index}/{total_legacy}")
                            last_status_time = now
            except Exception as exc:
                for rel_path, file_path, size, mtime_ns in legacy_doc_records:
                    current_files[rel_path] = self.extractor.build_error_record(
                        file_path=file_path,
                        size=size,
                        mtime_ns=mtime_ns,
                        message=f"Could not start shared Microsoft Word session: {exc}",
                    )

        deleted_count = len(set(previous_files.keys()) - set(current_files.keys()))

        cache_document = self.new_cache_document(root, allowed_extensions)
        cache_document["files"] = dict(sorted(current_files.items()))
        cache_document["manifest"] = {"directories": discovery_result.manifest_directories}

        previous_extensions = (
            sorted({ext.lower() for ext in previous_cache.get("allowed_extensions", [])})
            if previous_cache
            else []
        )
        current_extensions = sorted({ext.lower() for ext in allowed_extensions})
        previous_manifest = previous_cache.get("manifest", {}).get("directories", {}) if previous_cache else {}

        # `generated_at_utc` is intentionally excluded here. It is metadata about
        # the write operation itself, not a reason to rewrite an otherwise identical
        # cache document.
        cache_changed = (
            previous_cache is None
            or refreshed_count > 0
            or deleted_count > 0
            or previous_cache.get("include_subfolders") != root.include_subfolders
            or previous_extensions != current_extensions
            or len(previous_files) != len(current_files)
            or previous_manifest != discovery_result.manifest_directories
        )

        if cache_changed:
            self.write_cache(root.path, cache_document)
        else:
            cache_document = previous_cache

        status_callback(
            f"{root.label}: cache {'updated' if cache_changed else 'reused'} "
            f"(reused {reused_count}, refreshed {refreshed_count}, removed {deleted_count})"
        )
        return cache_document

    def search_record(
        self,
        root: SearchRoot,
        relative_path: str,
        record: Dict[str, object],
        search_terms: Sequence[str],
        match_mode: str,
    ) -> Iterator[Tuple[str, object]]:
        """
        Search one cached file record and yield UI-ready result payloads.
        """
        file_path = root.path / Path(relative_path)
        error_message = record.get("error")
        if error_message:
            yield ("text", format_error_block(error_message))
            return

        file_type = record.get("file_type")
        if file_type == "excel":
            # Excel matches are yielded one row at a time in v10.2 because the
            # text-output UI renders each hit as its own block.
            for row in record.get("rows", []):
                row_text = row.get("row_text", "")
                if not text_matches_search(row_text, search_terms, match_mode):
                    continue
                yield (
                    "result",
                    build_excel_result_payload(
                        file_path=file_path,
                        sheet=row.get("sheet"),
                        row_number=row.get("row_number"),
                        row_text=row_text,
                    ),
                )
            return

        if file_type in {"word", "pdf"}:
            # Text documents show the first matching line plus one line of context
            # before and after the match to make the hit easier to interpret.
            lines = record.get("lines", [])
            matching_indexes = [
                index for index, line in enumerate(lines)
                if text_matches_search(line, search_terms, match_mode)
            ]
            if not matching_indexes:
                return

            first_index = matching_indexes[0]
            line_before = lines[first_index - 1] if first_index > 0 else ""
            matched_line = lines[first_index]
            line_after = lines[first_index + 1] if first_index + 1 < len(lines) else ""

            yield (
                "result",
                build_text_result_payload(
                    file_path=file_path,
                    file_type=file_type,
                    line_number=first_index + 1,
                    line_before=line_before,
                    matched_line=matched_line,
                    line_after=line_after,
                    match_count=len(matching_indexes),
                ),
            )


# =============================================================================
# Safe external opening helper
# =============================================================================
class FileOpener:
    """Open result files as safely as the platform allows."""

    @staticmethod
    def open_result(file_type: str, file_path: str, sheet_name=None, row_number=None) -> None:
        """
        Open a selected result using the safest specialised opener available.
        """
        if file_type == "excel" and HAS_WIN32_COM and sheet_name and row_number and sys.platform.startswith("win"):
            FileOpener.open_excel_readonly(file_path=file_path, sheet_name=sheet_name, row_number=row_number)
            return

        if file_type == "word" and HAS_WIN32_COM and sys.platform.startswith("win"):
            FileOpener.open_word_readonly(file_path=file_path)
            return

        FileOpener.open_default(file_path)

    @staticmethod
    def open_default(file_path: str) -> None:
        """
        Fall back to the operating system's default file-open behaviour.
        """
        target_path = Path(file_path)
        if not target_path.exists():
            messagebox.showerror("File not found", f"The file does not exist:\n{file_path}")
            return

        try:
            if sys.platform.startswith("win"):
                os.startfile(str(target_path))
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(target_path)])
            else:
                subprocess.Popen(["xdg-open", str(target_path)])
        except Exception as exc:
            messagebox.showerror("Could not open file", str(exc))

    @staticmethod
    def open_word_readonly(file_path: str) -> None:
        """
        Open a Word document in Microsoft Word with read-only settings when possible.
        """
        target_path = Path(file_path)
        if not target_path.exists():
            messagebox.showerror("File not found", f"The document does not exist:\n{file_path}")
            return

        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = True
            word.DisplayAlerts = 0
            try:
                word.AutomationSecurity = 3
            except Exception:
                pass
            document = word.Documents.Open(
                str(target_path.resolve()),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=True,
                Revert=False,
                OpenAndRepair=False,
                NoEncodingDialog=True,
            )
            document.Activate()
        except Exception as exc:
            messagebox.showerror("Could not open document in Word", str(exc))
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    @staticmethod
    def open_excel_readonly(file_path: str, sheet_name: str, row_number: int) -> None:
        """
        Open an Excel workbook read-only and jump to the first matching sheet/row when possible.
        """
        target_path = Path(file_path)
        if not target_path.exists():
            messagebox.showerror("File not found", f"The workbook does not exist:\n{file_path}")
            return

        try:
            pythoncom.CoInitialize()
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = True
            try:
                excel.AutomationSecurity = 3
            except Exception:
                pass
            workbook = excel.Workbooks.Open(
                str(target_path.resolve()),
                UpdateLinks=0,
                ReadOnly=True,
                IgnoreReadOnlyRecommended=True,
                AddToMru=False,
                Notify=False,
            )
            worksheet = workbook.Worksheets(sheet_name)
            worksheet.Activate()
            target_cell = worksheet.Cells(row_number, 1)
            try:
                excel.Goto(target_cell, True)
            except Exception:
                target_cell.Select()
            try:
                excel.ActiveWindow.ScrollRow = max(1, row_number - 5)
                excel.ActiveWindow.ScrollColumn = 1
            except Exception:
                pass
            workbook.Activate()
        except Exception as exc:
            messagebox.showerror("Could not open workbook in Excel", str(exc))
        finally:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


# =============================================================================
# Search thread worker
# =============================================================================
class SearchRunner:
    """Coordinates discovery, cache refresh, and cache-based searching."""

    def __init__(
        self,
        cache_manager: CacheManager,
        queue_out: queue.Queue,
        cancel_event: threading.Event,
    ) -> None:
        """
        Store the shared cache manager, output queue, and cancellation event.
        """
        self.cache_manager = cache_manager
        self.queue = queue_out
        self.cancel_event = cancel_event

    def _status(self, message: str) -> None:
        """
        Send a lightweight status update back to the Tk main thread.
        """
        self.queue.put(("status", message))

    def _emit_items(self, items: List[Dict[str, object]]) -> None:
        """
        Flush a batch of queued UI items to reduce cross-thread chatter.
        """
        if items:
            self.queue.put(("items", items.copy()))
            items.clear()

    def run(self, roots: Sequence[SearchRoot], allowed_extensions: Sequence[str], search_terms: Sequence[str], match_mode: str) -> None:
        """
        Perform discovery, cache refresh, and cached searching on a worker thread.
        """
        started_at = time.perf_counter()
        files_scanned = 0
        matches_found = 0

        try:
            self._status("Preparing search...")
            cache_groups: List[Tuple[SearchRoot, Dict[str, object]]] = []
            total_files = 0

            # Discover files and refresh caches for each enabled root before the
            # actual content search begins.
            for root in roots:
                if self.cancel_event.is_set():
                    self.queue.put(("finished", {
                        "files_scanned": 0,
                        "matches_found": 0,
                        "cancelled": True,
                        "elapsed_seconds": time.perf_counter() - started_at,
                    }))
                    return

                previous_cache = self.cache_manager.load_cache(root.path)
                discovery_result = self.cache_manager.discover_root_files(
                    root=root,
                    allowed_extensions=allowed_extensions,
                    cancel_event=self.cancel_event,
                    status_callback=self._status,
                    previous_cache=previous_cache,
                )
                total_files += len(discovery_result.file_records)

                self._status(
                    f"{root.label}: found {len(discovery_result.file_records)} {pluralize(len(discovery_result.file_records), 'file')} "
                    f"({'manifest-assisted' if discovery_result.used_manifest else 'full scan'})"
                )

                cache_document = self.cache_manager.refresh_root_cache(
                    root=root,
                    discovery_result=discovery_result,
                    allowed_extensions=allowed_extensions,
                    previous_cache=previous_cache,
                    cancel_event=self.cancel_event,
                    status_callback=self._status,
                )
                cache_groups.append((root, cache_document))

            if total_files == 0:
                self.queue.put(("finished", {
                    "files_scanned": 0,
                    "matches_found": 0,
                    "cancelled": False,
                    "elapsed_seconds": time.perf_counter() - started_at,
                }))
                return

            item_buffer: List[Dict[str, object]] = []
            last_status_time = 0.0

            for root, cache_document in cache_groups:
                if self.cancel_event.is_set():
                    break

                files = cache_document.get("files", {})
                if not isinstance(files, dict):
                    continue

                # Search the cached content for every file belonging to this root.
                for rel_path, record in files.items():
                    if self.cancel_event.is_set():
                        break

                    files_scanned += 1
                    now = time.monotonic()
                    if files_scanned == 1 or files_scanned == total_files or (now - last_status_time) >= STATUS_UPDATE_INTERVAL_SECONDS:
                        self._status(f"Searching cached content {files_scanned}/{total_files}: {Path(rel_path).name}")
                        last_status_time = now

                    for item_type, payload in self.cache_manager.search_record(root, rel_path, record, search_terms, match_mode):
                        if self.cancel_event.is_set():
                            break

                        if item_type == "result":
                            matches_found += 1

                        item_buffer.append({"kind": item_type, "payload": payload})
                        if len(item_buffer) >= OUTPUT_BATCH_BLOCKS:
                            self._emit_items(item_buffer)

            self._emit_items(item_buffer)
            self.queue.put(("finished", {
                "files_scanned": files_scanned,
                "matches_found": matches_found,
                "cancelled": self.cancel_event.is_set(),
                "elapsed_seconds": time.perf_counter() - started_at,
            }))
        except Exception as exc:
            if self.cancel_event.is_set():
                self.queue.put(("finished", {
                    "files_scanned": files_scanned,
                    "matches_found": matches_found,
                    "cancelled": True,
                    "elapsed_seconds": time.perf_counter() - started_at,
                }))
            else:
                self.queue.put(("fatal", str(exc)))


# =============================================================================
# Tkinter UI
# =============================================================================
class ArchiveSearchApp:
    """Main Tkinter application."""

    def __init__(self, root: tk.Tk) -> None:
        """
        Create the main application state, Tk variables, and background-search plumbing.
        """
        self.root = root
        self.root.title("Archive Search")
        self.root.geometry("1350x820")

        self.app_folder = get_app_folder()
        self.queue: queue.Queue = queue.Queue()
        self.cancel_event = threading.Event()
        self.search_thread: Optional[threading.Thread] = None

        self.extractor = DocumentExtractor()
        self.cache_manager = CacheManager(self.extractor)

        self.current_highlight_terms: List[str] = []
        self.result_links: Dict[str, Dict[str, object]] = {}
        self.result_counter = 0

        self.primary_folder_var = tk.StringVar(value=str(DEFAULT_PRIMARY_SEARCH_FOLDER))
        self.secondary_folder_var = tk.StringVar(value=str(DEFAULT_SECONDARY_SEARCH_FOLDER))
        self.use_primary_var = tk.BooleanVar(value=True)
        self.use_secondary_var = tk.BooleanVar(value=True)

        self.primary_subfolders_var = tk.BooleanVar(value=True)
        self.secondary_subfolders_var = tk.BooleanVar(value=True)

        self.search_workbooks_var = tk.BooleanVar(value=True)
        self.search_text_documents_var = tk.BooleanVar(value=True)

        self.search_var = tk.StringVar()
        self.match_mode_var = tk.StringVar(value="any")
        self.status_var = tk.StringVar(value="Ready.")

        self._build_ui()
        self._update_folder_states()
        self._poll_queue()

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------
    def _build_ui(self) -> None:
        """
        Construct the full Tkinter user interface for the selected program version.
        """
        # Build the overall single-pane layout used by the v10.2 interface.
        container = ttk.Frame(self.root, padding=12)
        container.pack(fill="both", expand=True)

        top = ttk.Frame(container)
        top.pack(fill="x")

        style = ttk.Style(self.root)
        style.configure("FolderOn.TCheckbutton", font=("Segoe UI", 10, "bold"))
        style.configure("FolderOff.TCheckbutton", font=("Segoe UI", 10, "bold"), foreground="#808080")

        self.primary_check = ttk.Checkbutton(
            top,
            text="Primary Search Folder:",
            variable=self.use_primary_var,
            command=self._update_folder_states,
            style="FolderOn.TCheckbutton",
        )
        self.primary_check.grid(row=0, column=0, sticky="w", padx=(0, 4))

        self.primary_folder_entry = ttk.Entry(top, textvariable=self.primary_folder_var, width=90)
        self.primary_folder_entry.grid(row=0, column=1, sticky="we", padx=(4, 8))

        self.primary_browse_button = ttk.Button(
            top,
            text="Browse...",
            command=lambda: self._browse_folder(self.primary_folder_var, "Select primary search folder"),
        )
        self.primary_browse_button.grid(row=0, column=2, sticky="e")

        self.primary_subfolders_check = ttk.Checkbutton(
            top,
            text="Search subfolders",
            variable=self.primary_subfolders_var,
        )
        self.primary_subfolders_check.grid(row=0, column=3, sticky="w", padx=(8, 0))

        self.secondary_check = ttk.Checkbutton(
            top,
            text="Secondary Search Folder:",
            variable=self.use_secondary_var,
            command=self._update_folder_states,
            style="FolderOn.TCheckbutton",
        )
        self.secondary_check.grid(row=1, column=0, sticky="w", padx=(0, 4), pady=(8, 0))

        self.secondary_folder_entry = ttk.Entry(top, textvariable=self.secondary_folder_var, width=90)
        self.secondary_folder_entry.grid(row=1, column=1, sticky="we", padx=(4, 8), pady=(8, 0))

        self.secondary_browse_button = ttk.Button(
            top,
            text="Browse...",
            command=lambda: self._browse_folder(self.secondary_folder_var, "Select secondary search folder"),
        )
        self.secondary_browse_button.grid(row=1, column=2, sticky="e", pady=(8, 0))

        self.secondary_subfolders_check = ttk.Checkbutton(
            top,
            text="Search subfolders",
            variable=self.secondary_subfolders_var,
        )
        self.secondary_subfolders_check.grid(row=1, column=3, sticky="w", padx=(8, 0), pady=(8, 0))

        ttk.Label(top, text="Search Terms:", font=("Segoe UI", 10, "bold")).grid(
            row=2, column=0, sticky="w", pady=(12, 0)
        )
        self.search_entry = ttk.Entry(top, textvariable=self.search_var, width=70)
        self.search_entry.grid(row=2, column=1, columnspan=3, sticky="we", padx=(4, 0), pady=(12, 0))
        self.search_entry.bind("<Return>", lambda event: self.start_search())

        options = ttk.Frame(container)
        options.pack(fill="x", pady=(10, 0))

        ttk.Label(options, text="Match mode:").pack(side="left")
        ttk.Radiobutton(options, text="Any word", variable=self.match_mode_var, value="any").pack(side="left", padx=(8, 0))
        ttk.Radiobutton(options, text="All words", variable=self.match_mode_var, value="all").pack(side="left", padx=(8, 0))

        ttk.Label(options, text="File types:").pack(side="left", padx=(20, 0))
        self.workbooks_check = ttk.Checkbutton(options, text="Workbooks", variable=self.search_workbooks_var)
        self.workbooks_check.pack(side="left", padx=(8, 0))
        self.text_documents_check = ttk.Checkbutton(options, text="Text documents", variable=self.search_text_documents_var)
        self.text_documents_check.pack(side="left", padx=(8, 0))

        self.search_button = ttk.Button(options, text="Search", command=self.start_search)
        self.search_button.pack(side="left", padx=(20, 0))
        self.cancel_button = ttk.Button(options, text="Cancel Search", command=self.cancel_search, state="disabled")
        self.cancel_button.pack(side="left", padx=(8, 0))
        self.clear_button = ttk.Button(options, text="Clear", command=self.clear_output)
        self.clear_button.pack(side="left", padx=(8, 0))

        help_text = (
            "Select folders, enter search terms, then click Search.\n"
            "Use quotes for exact phrases, e.g. \"Sigma 5E\".\n"
            "Any word = any term | All words = every term.\n"
            "Click a result path to open the file 'read-only' where supported."
        )
        ttk.Label(container, text=help_text, justify="left").pack(anchor="w", pady=(10, 4))

        ttk.Label(container, textvariable=self.status_var).pack(anchor="w", pady=(0, 8))

        # The text widget doubles as both the result list and the detail view.
        text_frame = ttk.Frame(container)
        text_frame.pack(fill="both", expand=True)

        self.output = tk.Text(text_frame, wrap="word", font=("Consolas", 10))
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=self.output.yview)
        self.output.configure(yscrollcommand=scrollbar.set)

        self.output.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        self.output.tag_configure(HIGHLIGHT_TAG_NAME, background=HIGHLIGHT_BACKGROUND, foreground=HIGHLIGHT_FOREGROUND)
        self.output.tag_configure("sel", background="#cfe8ff", foreground="#000000")
        self.output.tag_configure("path_link", foreground="#0563C1", underline=True)
        self.output.tag_raise("sel")
        self.output.tag_lower(HIGHLIGHT_TAG_NAME, "sel")
        self.output.config(state="disabled")

        self.write_line("Archive Search")
        self.write_line("=" * 100)
        self.write_line(
            f"Primary folder  : {self.primary_folder_var.get()} "
            f"({'enabled' if self.use_primary_var.get() else 'disabled'}, "
            f"subfolders {'on' if self.primary_subfolders_var.get() else 'off'})"
        )
        self.write_line(
            f"Secondary folder: {self.secondary_folder_var.get()} "
            f"({'enabled' if self.use_secondary_var.get() else 'disabled'}, "
            f"subfolders {'on' if self.secondary_subfolders_var.get() else 'off'})"
        )
        self.write_line("File types      : Workbooks, Text documents")
        self.write_line("")

        top.columnconfigure(1, weight=1)

    # ------------------------------------------------------------------
    # UI state helpers
    # ------------------------------------------------------------------
    def _browse_folder(self, variable: tk.StringVar, title: str) -> None:
        """
        Open a folder picker and copy the user's selection into the bound Tk variable.
        """
        current = variable.get().strip()
        initial_dir = current if current and Path(current).exists() else str(self.app_folder)
        selected = filedialog.askdirectory(title=title, initialdir=initial_dir)
        if selected:
            variable.set(selected)

    def _update_folder_states(self) -> None:
        """
        Enable or disable folder controls based on which roots are active.
        """
        primary_enabled = self.use_primary_var.get()
        secondary_enabled = self.use_secondary_var.get()

        primary_state = "normal" if primary_enabled else "disabled"
        secondary_state = "normal" if secondary_enabled else "disabled"

        self.primary_folder_entry.config(state=primary_state)
        self.primary_browse_button.config(state=primary_state)
        self.primary_subfolders_check.config(state=primary_state)
        self.primary_check.config(
            style="FolderOn.TCheckbutton" if primary_enabled else "FolderOff.TCheckbutton"
        )

        self.secondary_folder_entry.config(state=secondary_state)
        self.secondary_browse_button.config(state=secondary_state)
        self.secondary_subfolders_check.config(state=secondary_state)
        self.secondary_check.config(
            style="FolderOn.TCheckbutton" if secondary_enabled else "FolderOff.TCheckbutton"
        )

    def _set_search_running_ui(self, running: bool) -> None:
        """
        Lock or unlock controls while a background search is running.
        """
        if running:
            self.search_button.config(state="disabled")
            self.cancel_button.config(state="normal")
            self.primary_check.config(state="disabled")
            self.secondary_check.config(state="disabled")
            self.primary_folder_entry.config(state="disabled")
            self.secondary_folder_entry.config(state="disabled")
            self.primary_browse_button.config(state="disabled")
            self.secondary_browse_button.config(state="disabled")
            self.primary_subfolders_check.config(state="disabled")
            self.secondary_subfolders_check.config(state="disabled")
            self.workbooks_check.config(state="disabled")
            self.text_documents_check.config(state="disabled")
            self.search_entry.config(state="disabled")
        else:
            self.search_button.config(state="normal")
            self.cancel_button.config(state="disabled")
            self.primary_check.config(state="normal")
            self.secondary_check.config(state="normal")
            self.workbooks_check.config(state="normal")
            self.text_documents_check.config(state="normal")
            self.search_entry.config(state="normal")
            self._update_folder_states()

    # ------------------------------------------------------------------
    # Output helpers
    # ------------------------------------------------------------------
    def _set_output_editable(self, editable: bool) -> None:
        """
        Temporarily unlock or relock the text output widget for updates.
        """
        self.output.config(state="normal" if editable else "disabled")

    def _insert_output_text(self, text: str, tags: Tuple[str, ...] = ()) -> None:
        """
        Insert text into the output area and highlight any active search terms.

        This helper no longer returns widget indexes because no caller needs them;
        it simply performs the insertion work in one place.
        """
        if not text:
            return

        start_index = self.output.index("end-1c")
        self.output.insert("end", text, tags)
        end_index = self.output.index("end-1c")

        if self.current_highlight_terms:
            self._highlight_inserted_range(start_index, end_index)


    def write_text(self, text: str) -> None:
        """
        Append raw text to the output pane and keep the newest content in view.
        """
        if not text:
            return
        self._set_output_editable(True)
        try:
            self._insert_output_text(text)
            self.output.see("end")
        finally:
            self._set_output_editable(False)

    def write_line(self, text: str = "") -> None:
        """
        Append one line of text to the output pane.
        """
        self.write_text(text + "\n")

    def clear_output(self) -> None:
        """
        Clear all rendered output and reset clickable result-link state.
        """
        self._set_output_editable(True)
        try:
            self.output.delete("1.0", "end")
            self.result_links.clear()
            self.result_counter = 0
        finally:
            self._set_output_editable(False)

    def _highlight_inserted_range(self, start_index: str, end_index: str) -> None:
        """
        Highlight search terms inside the most recently inserted output range.
        """
        for term in self.current_highlight_terms:
            if not term:
                continue

            search_from = start_index
            while True:
                match_start = self.output.search(pattern=term, index=search_from, stopindex=end_index, nocase=True)
                if not match_start:
                    break
                match_end = f"{match_start}+{len(term)}c"
                self.output.tag_add(HIGHLIGHT_TAG_NAME, match_start, match_end)
                search_from = match_end

    def write_result_block(self, result: SearchResult) -> None:
        """
        Render one search result as a formatted block in the scrolling text pane.
        """
        tag_name = f"path_link_{self.result_counter}"
        self.result_counter += 1

        self.result_links[tag_name] = {
            "file_type": result.file_type,
            "path": result.path,
            "sheet": result.sheet,
            "row_number": result.row_number,
        }

        self._set_output_editable(True)
        try:
            self._insert_output_text("=" * 100 + "\n\n")
            self._insert_output_text("Path: ")
            self._insert_output_text(result.path, ("path_link", tag_name))
            self._insert_output_text("\n\n")

            self.output.tag_bind(tag_name, "<Button-1>", self._on_result_click)
            self.output.tag_bind(tag_name, "<Enter>", lambda event: self.output.config(cursor="hand2"))
            self.output.tag_bind(tag_name, "<Leave>", lambda event: self.output.config(cursor="xterm"))

            if result.file_type == "excel":
                self._insert_output_text(
                    f"Document: {result.document_name}    Type: Excel    Sheet: {result.sheet}    Row: {result.row_number}\n\n"
                )
                self._insert_output_text("Content:\n")
                self._insert_output_text(f"  {result.row_text}\n\n")
            else:
                type_label = "Word" if result.file_type == "word" else "PDF"
                self._insert_output_text(f"Document: {result.document_name}    Type: {type_label}")
                if result.line_number is not None:
                    self._insert_output_text(f"    First line: {result.line_number}")
                if result.match_count is not None:
                    label = "mention" if result.match_count == 1 else "mentions"
                    self._insert_output_text(f"    Matches in file: {result.match_count} {label}")
                self._insert_output_text("\n\n")
                self._insert_output_text("First match context:\n")
                self._insert_output_text(f"{result.line_before or '—'}\n")
                self._insert_output_text(f"{result.matched_line or '—'}\n")
                self._insert_output_text(f"{result.line_after or '—'}\n\n")

            self.output.see("end")
        finally:
            self._set_output_editable(False)

    # ------------------------------------------------------------------
    # Result click handling
    # ------------------------------------------------------------------
    def _on_result_click(self, event) -> None:
        """
        Handle clicks on tagged file paths inside the text-output widget.
        """
        index = self.output.index(f"@{event.x},{event.y}")
        tags = self.output.tag_names(index)
        for tag in tags:
            if tag.startswith("path_link_") and tag in self.result_links:
                link = self.result_links[tag]
                FileOpener.open_result(
                    file_type=str(link["file_type"]),
                    file_path=str(link["path"]),
                    sheet_name=link.get("sheet"),
                    row_number=link.get("row_number"),
                )
                return

    # ------------------------------------------------------------------
    # Search setup
    # ------------------------------------------------------------------
    def _get_selected_search_roots(self) -> Optional[Tuple[SearchRoot, ...]]:
        """
        Validate the enabled folder selections and convert them into SearchRoot objects.
        """
        roots: List[SearchRoot] = []

        if not self.use_primary_var.get() and not self.use_secondary_var.get():
            messagebox.showwarning("No Search Folder Selected", "Please enable the primary folder, the secondary folder, or both.")
            return None

        if self.use_primary_var.get():
            primary_text = self.primary_folder_var.get().strip()
            if not primary_text:
                messagebox.showwarning("Missing Folder", "The primary folder is enabled, but no primary folder was entered.")
                return None
            primary_path = Path(primary_text).expanduser()
            if not primary_path.exists() or not primary_path.is_dir():
                messagebox.showerror(
                    "Invalid Primary Folder",
                    f"The primary search folder does not exist or is not a folder:\n{primary_path}",
                )
                return None
            roots.append(SearchRoot("Primary", primary_path, self.primary_subfolders_var.get()))

        if self.use_secondary_var.get():
            secondary_text = self.secondary_folder_var.get().strip()
            if not secondary_text:
                messagebox.showwarning("Missing Folder", "The secondary folder is enabled, but no secondary folder was entered.")
                return None
            secondary_path = Path(secondary_text).expanduser()
            if not secondary_path.exists() or not secondary_path.is_dir():
                messagebox.showerror(
                    "Invalid Secondary Folder",
                    f"The secondary search folder does not exist or is not a folder:\n{secondary_path}",
                )
                return None
            roots.append(SearchRoot("Secondary", secondary_path, self.secondary_subfolders_var.get()))

        return tuple(roots)

    def _get_allowed_extensions(self) -> Optional[Tuple[str, ...]]:
        """
        Translate the selected file-type checkboxes into an extension whitelist.
        """
        allowed_extensions = set()
        if self.search_workbooks_var.get():
            allowed_extensions |= WORKBOOK_EXTENSIONS
        if self.search_text_documents_var.get():
            allowed_extensions |= TEXT_DOCUMENT_EXTENSIONS
        if not allowed_extensions:
            messagebox.showwarning("No File Types Selected", "Please tick Workbooks, Text documents, or both.")
            return None
        return tuple(sorted(allowed_extensions))

    def _get_selected_file_type_label(self) -> str:
        """
        Return a human-readable label for the active file-type filter.
        """
        labels: List[str] = []
        if self.search_workbooks_var.get():
            labels.append("Workbooks")
        if self.search_text_documents_var.get():
            labels.append("Text documents")
        return ", ".join(labels) if labels else "None"

    def start_search(self) -> None:
        """
        Validate the UI state, prepare the worker thread, and start a new search.
        """
        if self.search_thread and self.search_thread.is_alive():
            return

        # Collect and validate all user selections before any background work starts.
        search_input = self.search_var.get().strip()
        if not search_input:
            messagebox.showwarning("Missing Search Terms", "Please type one or more search terms.")
            return

        try:
            search_terms = tuple(term.lower() for term in shlex.split(search_input))
        except ValueError as exc:
            messagebox.showerror("Invalid Search Input", f"Could not parse search terms:\n{exc}")
            return

        search_roots = self._get_selected_search_roots()
        if not search_roots:
            return
        allowed_extensions = self._get_allowed_extensions()
        if not allowed_extensions:
            return

        self.cancel_event.clear()

        # Highlights mirror the parsed search tokens rather than the raw input so
        # quoted phrases and individual words are handled consistently.
        self.current_highlight_terms = build_highlight_terms(search_terms)

        self.clear_output()
        self.write_line("Archive Search")
        self.write_line("=" * 100)

        primary_root = next((root for root in search_roots if root.label == "Primary"), None)
        secondary_root = next((root for root in search_roots if root.label == "Secondary"), None)

        if primary_root is not None:
            self.write_line(f"Primary folder  : {primary_root.path} (enabled, subfolders {'on' if primary_root.include_subfolders else 'off'})")
        else:
            self.write_line(f"Primary folder  : {self.primary_folder_var.get()} (disabled, subfolders {'on' if self.primary_subfolders_var.get() else 'off'})")

        if secondary_root is not None:
            self.write_line(f"Secondary folder: {secondary_root.path} (enabled, subfolders {'on' if secondary_root.include_subfolders else 'off'})")
        else:
            self.write_line(f"Secondary folder: {self.secondary_folder_var.get()} (disabled, subfolders {'on' if self.secondary_subfolders_var.get() else 'off'})")

        self.write_line(f"File types      : {self._get_selected_file_type_label()}")
        self.write_line(f"Search terms    : {search_input}")
        self.write_line(f"Match mode      : {self.match_mode_var.get()}")
        self.write_line("")

        self.status_var.set("Searching...")
        self._set_search_running_ui(True)

        runner = SearchRunner(
            cache_manager=self.cache_manager,
            queue_out=self.queue,
            cancel_event=self.cancel_event,
        )
        self.search_thread = threading.Thread(
            target=runner.run,
            args=(search_roots, allowed_extensions, search_terms, self.match_mode_var.get()),
            daemon=True,
        )
        self.search_thread.start()

    def cancel_search(self) -> None:
        """
        Signal the worker thread to stop at the next safe cancellation point.
        """
        if self.search_thread and self.search_thread.is_alive():
            self.cancel_event.set()
            self.status_var.set("Cancelling...")
            self.cancel_button.config(state="disabled")

    # ------------------------------------------------------------------
    # Queue polling
    # ------------------------------------------------------------------
    def _poll_queue(self) -> None:
        """
        Consume queued worker messages and reflect them safely in the Tk interface.
        """
        pending_item_batches: List[List[Dict[str, object]]] = []
        latest_status: Optional[str] = None
        finished_payload: Optional[Dict[str, object]] = None
        fatal_payload: Optional[str] = None

        try:
            # Drain everything currently queued so the UI updates in larger, smoother
            # batches instead of repainting for every single worker message.
            while True:
                item_type, payload = self.queue.get_nowait()
                if item_type == "items":
                    pending_item_batches.append(payload)
                elif item_type == "status":
                    latest_status = payload
                elif item_type == "finished":
                    finished_payload = payload
                elif item_type == "fatal":
                    fatal_payload = payload
        except queue.Empty:
            pass

        if pending_item_batches:
            for batch in pending_item_batches:
                for item in batch:
                    kind = item["kind"]
                    payload = item["payload"]
                    if kind == "text":
                        self.write_text(str(payload))
                    elif kind == "result":
                        self.write_result_block(payload)

        if latest_status is not None:
            self.status_var.set(latest_status)

        if finished_payload is not None:
            files_scanned = finished_payload["files_scanned"]
            matches_found = finished_payload["matches_found"]
            cancelled = finished_payload["cancelled"]

            scanned_label = pluralize(files_scanned, "file")
            result_label = pluralize(matches_found, "result")
            elapsed_seconds = finished_payload.get("elapsed_seconds", 0.0)

            if cancelled:
                final_message = (
                    f"Cancelled. Scanned {files_scanned} {scanned_label}. "
                    f"Found {matches_found} matching {result_label} before cancellation."
                    f"Time: {elapsed_seconds:.2f} seconds."
                )
            else:
                final_message = (
                    f"Finished. Scanned {files_scanned} {scanned_label}. "
                    f"Found {matches_found} matching {result_label}."
                    f"Time: {elapsed_seconds:.2f} seconds."
                )

            self.write_line("=" * 100)
            self.write_line(final_message)
            self.status_var.set(final_message)
            self._set_search_running_ui(False)

        if fatal_payload is not None:
            self.write_line(f"Fatal error: {fatal_payload}")
            self.status_var.set("Search failed.")
            self._set_search_running_ui(False)
            messagebox.showerror("Search failed", fatal_payload)

        self.root.after(100, self._poll_queue)


# =============================================================================
# Entrypoint
# =============================================================================
def main() -> None:
    """
    Create the Tk root window, start the application, and enter the Tk event loop.
    """
    root = tk.Tk()

    # The application object manages itself via Tk callbacks, so there is no need
    # to keep a separate local reference here.
    ArchiveSearchApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
